"""
DOCX Processor module for handling Word documents
"""

import os
from pathlib import Path
from typing import List, Dict, Optional, Union
from loguru import logger

import docx
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


class DocxProcessor:
    """
    Class for processing DOCX documents
    """
    
    def extract_text(self, docx_path: Union[str, Path]) -> str:
        """
        Extract text from a DOCX document
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            str: Extracted text
        """
        try:
            # Open the document
            doc = docx.Document(str(docx_path))
            
            # Extract text from paragraphs
            full_text = []
            
            # Process all elements (paragraphs and tables)
            for element in self._iter_block_items(doc):
                if isinstance(element, Paragraph):
                    # Extract text from paragraph
                    text = element.text.strip()
                    if text:
                        full_text.append(text)
                elif isinstance(element, Table):
                    # Extract text from table
                    table_text = self._extract_table_text(element)
                    if table_text:
                        full_text.append(table_text)
            
            # Join all text with double newlines
            return "\n\n".join(full_text)
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def _iter_block_items(self, doc: DocxDocument):
        """
        Iterate through all block items in the document (paragraphs and tables)
        
        Args:
            doc: DOCX document
            
        Yields:
            Union[Paragraph, Table]: Paragraph or Table object
        """
        # Get the document element
        body = doc._body._body
        
        # Iterate through all elements
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, doc._body)
            elif isinstance(child, CT_Tbl):
                yield Table(child, doc._body)
    
    def _extract_table_text(self, table: Table) -> str:
        """
        Extract text from a table
        
        Args:
            table: Table object
            
        Returns:
            str: Extracted text from the table
        """
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Get text from each cell
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                cells.append(cell_text)
            
            # Join cells with tabs
            row_text = "\t".join(cells)
            if row_text.strip():
                rows.append(row_text)
        
        # Join rows with newlines
        return "\n".join(rows)
    
    def extract_metadata(self, docx_path: Union[str, Path]) -> Dict[str, str]:
        """
        Extract metadata from a DOCX document
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dict[str, str]: Document metadata
        """
        try:
            # Open the document
            doc = docx.Document(str(docx_path))
            
            # Extract core properties
            core_props = doc.core_properties
            
            # Create metadata dictionary
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": str(core_props.revision) if core_props.revision else "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from DOCX: {str(e)}")
            return {}
    
    def extract_images(self, docx_path: Union[str, Path], output_dir: Path) -> List[Path]:
        """
        Extract images from a DOCX document
        
        Args:
            docx_path: Path to the DOCX file
            output_dir: Directory to save the images
            
        Returns:
            List[Path]: List of paths to the extracted images
        """
        try:
            # Create output directory if it doesn't exist
            os.makedirs(output_dir, exist_ok=True)
            
            # Open the document
            doc = docx.Document(str(docx_path))
            
            image_paths = []
            image_count = 0
            
            # Extract images from the document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    # Get image data
                    image_data = rel.target_part.blob
                    
                    # Determine image extension
                    image_ext = rel.target_ref.split(".")[-1]
                    
                    # Save image
                    image_path = output_dir / f"image_{image_count + 1}.{image_ext}"
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_data)
                    
                    image_paths.append(image_path)
                    image_count += 1
            
            logger.info(f"Extracted {image_count} images from DOCX")
            return image_paths
            
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {str(e)}")
            return []
    
    def extract_headers_footers(self, docx_path: Union[str, Path]) -> Dict[str, str]:
        """
        Extract headers and footers from a DOCX document
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dict[str, str]: Headers and footers
        """
        try:
            # Open the document
            doc = docx.Document(str(docx_path))
            
            headers = []
            footers = []
            
            # Extract headers
            for section in doc.sections:
                # Get header text
                if section.header.is_linked_to_previous:
                    continue
                
                header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
                if header_text:
                    headers.append(header_text)
                
                # Get footer text
                if section.footer.is_linked_to_previous:
                    continue
                
                footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
                if footer_text:
                    footers.append(footer_text)
            
            return {
                "headers": "\n".join(headers),
                "footers": "\n".join(footers)
            }
            
        except Exception as e:
            logger.error(f"Error extracting headers and footers from DOCX: {str(e)}")
            return {"headers": "", "footers": ""}
    
    def extract_structure(self, docx_path: Union[str, Path]) -> List[Dict]:
        """
        Extract document structure (headings and their levels)
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List[Dict]: Document structure
        """
        try:
            # Open the document
            doc = docx.Document(str(docx_path))
            
            structure = []
            
            # Extract headings
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    # Get heading level
                    level = int(paragraph.style.name.replace('Heading ', ''))
                    
                    # Add to structure
                    structure.append({
                        "level": level,
                        "text": paragraph.text.strip()
                    })
            
            return structure
            
        except Exception as e:
            logger.error(f"Error extracting structure from DOCX: {str(e)}")
            return []
